#reading
#basic reading
from docx import Document

# Load a DOCX file
doc = Document("example.docx")

# Loop through all paragraphs
for para in doc.paragraphs:
    print(para.text)

#using docx2txt
import docx2txt

text = docx2txt.process("example.docx")
print(text)

#using pandas
import pandas as pd

table = doc.tables[0]
data = [[cell.text for cell in row.cells] for row in table.rows]
df = pd.DataFrame(data[1:], columns=data[0])  # first row as header
print(df)

#write
#DOCX File
from docx import Document

# Create a new document
doc = Document()

# Add a heading
doc.add_heading("My First Document", level=1)

# Add a paragraph
doc.add_paragraph("This is my first paragraph in the document.")

# Save the document
doc.save("example.docx")

#visualising
from docx import Document
import pandas as pd
import matplotlib.pyplot as plt

# Load DOCX
doc = Document("example.docx")

# Select the first table
table = doc.tables[0]

# Extract table data
data = [[cell.text for cell in row.cells] for row in table.rows]
df = pd.DataFrame(data[1:], columns=data[0])  # first row as header

# Example: Plot a bar chart
df['Age'] = df['Age'].astype(int)  # convert column to int if numeric
df.plot(kind='bar', x='Name', y='Age')
plt.show()
